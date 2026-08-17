from pathlib import Path

from docx import Document

from app.models.extracted_document import (
    DocumentSection,
    ExtractedDocument,
)


class DOCXExtractor:

    @staticmethod
    def extract(
        file_path: str,
    ) -> ExtractedDocument:

        path = Path(file_path)

        document = Document(path)

        sections: list[DocumentSection] = []

        current_lines: list[str] = []
        section_number = 1

        full_text_parts: list[str] = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            full_text_parts.append(text)

            style_name = (
                paragraph.style.name
                if paragraph.style
                else ""
            )

            is_heading = (
                style_name.lower().startswith(
                    "heading"
                )
            )

            if (
                is_heading
                and current_lines
            ):
                sections.append(
                    DocumentSection(
                        section_number=section_number,
                        text="\n".join(
                            current_lines
                        ).strip(),
                    )
                )

                section_number += 1
                current_lines = []

            current_lines.append(text)

        if current_lines:
            sections.append(
                DocumentSection(
                    section_number=section_number,
                    text="\n".join(
                        current_lines
                    ).strip(),
                )
            )

        full_text = "\n\n".join(
            full_text_parts
        )

        return ExtractedDocument(
            filename=path.name,
            file_type="docx",
            text=full_text,
            sections=sections,
            page_count=None,
        )